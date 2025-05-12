import os
import argparse
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
import json
import time
import re
from docx.oxml.ns import qn
import textwrap

def extract_text_from_pdf(pdf_path):
    """从PDF文件中提取文本"""
    result = {
        'text_blocks': []
    }
    
    # 使用pdfplumber提取文本
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    result['text_blocks'].append({
                        'page': page_num,
                        'content': page_text
                    })
    except Exception as e:
        print(f"文本提取错误: {e}")
    
    # 按页码排序
    result['text_blocks'] = sorted(result['text_blocks'], key=lambda x: x['page'])
    
    return result

def split_into_paragraphs(text):
    """将文本分割成段落"""
    paragraphs = []
    current = []
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            if current:
                paragraphs.append(' '.join(current))
                current = []
        else:
            current.append(line)
    
    if current:
        paragraphs.append(' '.join(current))
    
    return paragraphs

def identify_figure_references(text):
    """识别文本中的图表引用"""
    # 常见的图表引用模式
    patterns = [
        r'Fig(?:ure)?\s*\.?\s*(\d+[a-zA-Z]*)',
        r'Figure\s+(\d+[a-zA-Z]*)',
        r'(Figure\s*\d+[a-zA-Z]*)',
        r'(Fig\.\s*\d+[a-zA-Z]*)'
    ]
    
    references = []
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            references.append((match.group(), match.start(), match.end()))
    
    return sorted(references, key=lambda x: x[1])

def translate_text_ollama(text, model="llama3", from_lang="en", to_lang="zh", max_retries=3):
    """使用Ollama本地模型翻译文本，添加重试机制"""
    # 如果文本过长，则分段翻译
    if len(text) > 1000:
        chunks = textwrap.wrap(text, 1000, break_long_words=False, replace_whitespace=False)
        translated_chunks = []
        
        for chunk in chunks:
            chunk_translation = translate_single_chunk_ollama(chunk, model, from_lang, to_lang, max_retries)
            translated_chunks.append(chunk_translation)
        
        return ' '.join(translated_chunks)
    else:
        return translate_single_chunk_ollama(text, model, from_lang, to_lang, max_retries)

def translate_single_chunk_ollama(text, model="llama3", from_lang="en", to_lang="zh", max_retries=3):
    """翻译单个文本块，包含重试机制"""
    retry_count = 0
    backoff_time = 1  # 初始等待时间(秒)
    
    while retry_count < max_retries:
        try:
            url = "http://localhost:11434/api/generate"
            
            # 构建更明确的翻译提示
            prompt = f"""请将以下英文文本精确、完整地翻译成标准学术中文。务必保持专业性，翻译内容不要有任何遗漏：

{text}

只返回中文翻译结果，不要包含原文或额外说明。"""
            
            payload = {
                "model": model,
                "prompt": prompt,
                "stream": False
            }
            
            response = requests.post(url, json=payload, timeout=60)
            if response.status_code == 200:
                result = response.json()
                translation = result.get("response", "").strip()
                
                # 验证翻译结果是否有效
                if len(translation) < len(text) * 0.3:  # 如果翻译结果太短，可能有问题
                    # print(f"警告：翻译结果过短，可能不完整")
                    if retry_count < max_retries - 1:
                        retry_count += 1
                        time.sleep(backoff_time)
                        backoff_time *= 2  # 指数退避
                        continue
                
                return translation
            else:
                print(f"Ollama API错误: {response.status_code}")
                if retry_count < max_retries - 1:
                    retry_count += 1
                    time.sleep(backoff_time)
                    backoff_time *= 2
                    continue
                return f"[Ollama翻译错误: 状态码 {response.status_code}]"
                
        except Exception as e:
            print(f"Ollama翻译错误: {e}")
            if retry_count < max_retries - 1:
                retry_count += 1
                time.sleep(backoff_time)
                backoff_time *= 2
                continue
            return f"[Ollama翻译错误: {text}]"
    
    # 所有重试都失败
    return f"[多次翻译失败，请手动翻译]"

def translate_text_google(text, from_lang="en", to_lang="zh-CN", max_retries=3):
    """使用Google翻译API翻译文本，包含分段和重试机制"""
    # 分段机制：Google翻译API对长文本有限制
    if len(text) > 5000:
        chunks = textwrap.wrap(text, 4500, break_long_words=False, replace_whitespace=False)
        translated_chunks = []
        
        for chunk in chunks:
            chunk_translation = translate_single_chunk_google(chunk, from_lang, to_lang, max_retries)
            translated_chunks.append(chunk_translation)
            # 避免翻译服务限制
            time.sleep(1)
        
        return '\n'.join(translated_chunks)
    else:
        return translate_single_chunk_google(text, from_lang, to_lang, max_retries)

def translate_single_chunk_google(text, from_lang="en", to_lang="zh-CN", max_retries=3):
    """翻译单个Google文本块，包含重试机制"""
    retry_count = 0
    backoff_time = 1  # 初始等待时间(秒)
    
    while retry_count < max_retries:
        try:
            url = "https://translate.googleapis.com/translate_a/single"
            
            params = {
                "client": "gtx",
                "sl": from_lang,
                "tl": to_lang,
                "dt": "t",
                "q": text
            }
            
            # 添加延迟避免请求过快
            time.sleep(0.5)
            
            response = requests.get(url, params=params, timeout=30)
            if response.status_code == 200:
                result = response.json()
                # 提取翻译结果并合并
                translated_text = ''.join([sentence[0] for sentence in result[0] if sentence[0]])
                
                # 验证翻译结果是否有效
                if len(translated_text) < len(text) * 0.3:  # 如果翻译结果太短，可能有问题
                    print(f"警告：Google翻译结果过短，可能不完整")
                    if retry_count < max_retries - 1:
                        retry_count += 1
                        time.sleep(backoff_time)
                        backoff_time *= 2
                        continue
                
                return translated_text.strip()
            else:
                print(f"Google翻译API错误: {response.status_code}")
                if retry_count < max_retries - 1:
                    retry_count += 1
                    time.sleep(backoff_time)
                    backoff_time *= 2
                    continue
                return f"[Google翻译错误: 状态码 {response.status_code}]"
                
        except Exception as e:
            print(f"Google翻译错误: {e}")
            if retry_count < max_retries - 1:
                retry_count += 1
                time.sleep(backoff_time)
                backoff_time *= 2
                continue
            return f"[Google翻译错误: {e}]"
    
    # 所有重试都失败
    return f"[多次翻译失败，请手动翻译]"

def apply_font_style(run, is_chinese=False):
    """应用字体样式，区分中英文"""
    # 小四字体大小（14磅）
    run.font.size = Pt(12)
    
    if is_chinese:
        # 中文使用宋体（小四）
        run.font.name = 'Times New Roman'  # 设置英文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    else:
        # 英文使用Times New Roman（小四）
        run.font.name = 'Times New Roman'

def create_bilingual_doc(pdf_data, output_path, translate_func=None):
    """创建中英文对照Word文档"""
    doc = Document()
    
    # 设置文档默认字体样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)  # 设置小四字体（12磅）
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置东亚文字字体为宋体
    
    # 设置文档标题
    title = doc.add_heading('外文文献翻译', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 为标题设置字体
    for run in title.runs:
        apply_font_style(run, is_chinese=True)
    
    # 按页码处理文本
    for text_block in pdf_data['text_blocks']:
        page_num = text_block['page']
        
        # 添加页码标记
        page_heading = doc.add_heading(f'第 {page_num} 页', level=2)
        page_heading.style.font.color.rgb = RGBColor(0, 0, 128)
        # 为页码标题设置字体
        for run in page_heading.runs:
            apply_font_style(run, is_chinese=True)
        
        # 分割成段落
        paragraphs = split_into_paragraphs(text_block['content'])
        
        for i, para_text in enumerate(paragraphs):
            # 跳过非常短的段落(可能是页码等)
            if len(para_text.strip()) < 10:
                continue
                
            # 查找图表引用
            figure_refs = identify_figure_references(para_text)
                
            # 添加原文段落
            para_num = i + 1
            para_heading = doc.add_heading(f'段落 {para_num}', level=3)
            # 为段落标题设置字体
            for run in para_heading.runs:
                apply_font_style(run, is_chinese=True)
            
            orig_para = doc.add_paragraph()
            orig_run = orig_para.add_run('原文：')
            orig_run.bold = True
            apply_font_style(orig_run, is_chinese=True)
            
            # 如果有图表引用，突出显示
            if figure_refs:
                last_pos = 0
                for ref, start, end in figure_refs:
                    # 添加引用前的文本
                    if start > last_pos:
                        run = orig_para.add_run(para_text[last_pos:start])
                        apply_font_style(run, is_chinese=False)
                    
                    # 添加高亮的引用
                    ref_run = orig_para.add_run(para_text[start:end])
                    ref_run.bold = True
                    ref_run.font.color.rgb = RGBColor(255, 0, 0)
                    apply_font_style(ref_run, is_chinese=False)
                    
                    last_pos = end
                
                # 添加最后一个引用后的文本
                if last_pos < len(para_text):
                    run = orig_para.add_run(para_text[last_pos:])
                    apply_font_style(run, is_chinese=False)
                
                # 添加图片位置标记
                img_placeholder = doc.add_paragraph()
                placeholder_run = img_placeholder.add_run('【图片位置】')
                placeholder_run.bold = True
                apply_font_style(placeholder_run, is_chinese=True)
                
                img_desc_run = img_placeholder.add_run(' 请在此处插入对应的图片和说明')
                apply_font_style(img_desc_run, is_chinese=True)
                img_placeholder.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                run = orig_para.add_run(para_text)
                apply_font_style(run, is_chinese=False)
            
            # 添加翻译段落
            trans_para = doc.add_paragraph()
            trans_label = trans_para.add_run('翻译：')
            trans_label.bold = True
            apply_font_style(trans_label, is_chinese=True)
            
            if translate_func:
                # 执行翻译
                print(f"正在翻译第 {page_num} 页 段落 {para_num}...")
                translation = translate_func(para_text)
                
                # 设置中文字体
                trans_run = trans_para.add_run(translation)
                apply_font_style(trans_run, is_chinese=True)
            else:
                # 没有翻译功能，添加占位符
                placeholder = trans_para.add_run('[请在此处添加中文翻译]')
                apply_font_style(placeholder, is_chinese=True)
            
            # 添加分隔线
            separator = doc.add_paragraph('_' * 50)
            sep_run = separator.runs[0]
            apply_font_style(sep_run, is_chinese=False)
    
    # 保存文档
    doc.save(output_path)
    print(f"已创建文档: {output_path}")

def main():
    parser = argparse.ArgumentParser(description='PDF文献翻译工具')
    parser.add_argument('pdf_path', help='PDF文件路径')
    parser.add_argument('--output', '-o', help='输出Word文件路径')
    parser.add_argument('--translate', '-t', choices=['none', 'ollama', 'google'], default='none', 
                      help='翻译方式: none(无翻译), ollama(本地Ollama), google(Google翻译)')
    parser.add_argument('--model', '-m', default='llama3', help='Ollama翻译使用的模型名称')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.pdf_path):
        print(f"错误: 找不到PDF文件 '{args.pdf_path}'")
        return
    
    # 设置输出路径
    if not args.output:
        pdf_name = os.path.splitext(os.path.basename(args.pdf_path))[0]
        args.output = f"{pdf_name}_translation.docx"
    
    # 提取PDF文本
    print(f"正在处理PDF: {args.pdf_path}")
    pdf_data = extract_text_from_pdf(args.pdf_path)
    
    if not pdf_data['text_blocks']:
        print("错误: 无法提取PDF文本")
        return
    
    print(f"成功提取 {len(pdf_data['text_blocks'])} 页文本")
    
    # 检查Ollama服务（如果需要）
    if args.translate == 'ollama':
        try:
            test_response = requests.get("http://localhost:11434/api/tags", timeout=5)
            if test_response.status_code != 200:
                print("警告: Ollama服务可能未正常运行。建议使用Google翻译('-t google')。")
                choice = input("是否切换到Google翻译? (y/n): ").strip().lower()
                if choice == 'y':
                    args.translate = 'google'
        except:
            print("警告: 无法连接到Ollama服务。建议使用Google翻译('-t google')。")
            choice = input("是否切换到Google翻译? (y/n): ").strip().lower()
            if choice == 'y':
                args.translate = 'google'
    
    # 创建Word文档
    if args.translate == 'ollama':
        print(f"正在使用Ollama({args.model})创建带翻译的Word文档...")
        # 为当前会话创建一个带有选定模型的翻译函数
        translate_func = lambda text: translate_text_ollama(text, model=args.model)
        create_bilingual_doc(pdf_data, args.output, translate_func)
    elif args.translate == 'google':
        print("正在使用Google翻译创建带翻译的Word文档...")
        translate_func = translate_text_google
        create_bilingual_doc(pdf_data, args.output, translate_func)
    else:
        print("正在创建Word文档(无自动翻译)...")
        create_bilingual_doc(pdf_data, args.output)
    
    print(f"完成! 文档已保存至: {args.output}")

if __name__ == "__main__":
    main() 